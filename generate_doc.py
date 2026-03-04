"""
Generates PROJECT_DOCUMENTATION.docx — a comprehensive Word document for the
PDE Claim Processing System (Dynamic Multi-Agent AI).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_page_margins(document, top=1.0, bottom=1.0, left=1.2, right=1.2):
    section = document.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)

def rgb(r, g, b):
    return RGBColor(r, g, b)

def add_heading(document, text, level=1):
    p = document.add_heading(text, level=level)
    if level == 1:
        for run in p.runs:
            run.font.color.rgb = rgb(0, 70, 127)
            run.font.size = Pt(18)
    elif level == 2:
        for run in p.runs:
            run.font.color.rgb = rgb(0, 102, 178)
            run.font.size = Pt(14)
    elif level == 3:
        for run in p.runs:
            run.font.color.rgb = rgb(31, 73, 125)
            run.font.size = Pt(12)
    return p

def add_para(document, text="", bold=False, italic=False, size=10.5, color=None):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(document, text, level=0, bold_prefix=None):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.25 * level)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10.5)
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
    return p

def add_code_block(document, text):
    """Add a monospaced code / diagram block with grey background."""
    p = document.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb(40, 40, 40)
    return p

def shade_row(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

def add_table(document, headers, rows, col_widths=None, header_color="1F4E79"):
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    shade_row(hdr, header_color)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = rgb(255, 255, 255)
        run.font.size = Pt(10)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        shade_row(row, "EBF3FB" if ri % 2 == 0 else "FFFFFF")
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_info_box(document, text, bg="E8F4FD", border_color="2196F3"):
    """Add a shaded paragraph acting as an info/note box."""
    p = document.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), bg)
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = rgb(13, 71, 161)
    return p

def page_break(document):
    document.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# Setup
# ─────────────────────────────────────────────────────────────────────────────
set_page_margins(doc)

# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
doc.add_paragraph("\n\n\n")

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
t_run = title_p.add_run("💊  PDE Claim Processing System")
t_run.bold = True
t_run.font.size = Pt(26)
t_run.font.color.rgb = rgb(0, 70, 127)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
s_run = sub_p.add_run("Dynamic Multi-Agent AI — Complete Workflow & Architecture")
s_run.font.size = Pt(14)
s_run.font.color.rgb = rgb(0, 102, 178)
s_run.italic = True

doc.add_paragraph("\n")

tags_p = doc.add_paragraph()
tags_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tags_run = tags_p.add_run(
    "LangGraph  ·  Azure OpenAI GPT-4o  ·  ReAct Orchestrator  ·  LangChain Tools  ·  Streamlit  ·  SQLite"
)
tags_run.font.size = Pt(10.5)
tags_run.font.color.rgb = rgb(80, 80, 80)

doc.add_paragraph("\n\n")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Date: March 2026    |    Medicare Part D — PDE Error Code Processing\n").font.size = Pt(10)

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# TABLE OF CONTENTS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Table of Contents", level=1)

toc_items = [
    ("1.", "Project Overview"),
    ("2.", "System Architecture"),
    ("3.", "Dynamic Orchestration Flow"),
    ("4.", "Agent Descriptions"),
    ("5.", "Decision Logic & Routing Table"),
    ("6.", "Scenario Walkthroughs"),
    ("7.", "Database Schema"),
    ("8.", "Tools Reference"),
    ("9.", "File Structure"),
    ("10.", "Setup & Running the Application"),
    ("11.", "Testing"),
    ("12.", "Key Design Principles"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r1 = p.add_run(f"  {num}  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = rgb(0, 70, 127)
    r2 = p.add_run(title)
    r2.font.size = Pt(11)

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 1. PROJECT OVERVIEW
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "1.  Project Overview", level=1)

add_para(doc,
    "The PDE Claim Processing System is a fully dynamic, LLM-driven multi-agent AI pipeline "
    "that processes Medicare Part D Prescription Drug Event (PDE) claims flagged by CMS "
    "(Centers for Medicare & Medicaid Services) with specific error codes.",
    size=11)

doc.add_paragraph()

add_heading(doc, "What the System Does", level=2)

steps = [
    "Reads the relevant Standard Operating Procedure (SOP) document for the given CMS error code.",
    "Fetches the claim record from the SQLite database.",
    "Uses an LLM-powered ReAct Orchestrator to decide at runtime which specialist agents to invoke and in what order.",
    "Calls each required specialist agent as a LangChain tool — each agent handles one specific responsibility.",
    "Collects all agent outputs into a unified LangGraph state.",
    "Displays full results, reasoning traces, and reports in a Streamlit dashboard.",
]
for i, s in enumerate(steps, 1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run(s)
    r.font.size = Pt(10.5)

doc.add_paragraph()
add_heading(doc, "Supported CMS Error Codes", level=2)

add_table(doc,
    headers=["Error Code", "Description", "Possible Outcomes"],
    rows=[
        ["781", "Provider ID missing or does not exist in CMS master provider file", "REPROCESS  or  REJECT"],
        ["935", "Claim has already been adjudicated and submitted previously to CMS", "REPROCESS  or  ALREADY_PROCESSED"],
    ],
    col_widths=[1.2, 3.5, 2.2],
)

doc.add_paragraph()
add_heading(doc, "Technology Stack", level=2)

add_table(doc,
    headers=["Component", "Technology"],
    rows=[
        ["LLM / AI Engine", "Azure OpenAI GPT-4o (AzureChatOpenAI)"],
        ["Agent Framework", "LangChain  create_agent()  with ReAct loop"],
        ["Pipeline Orchestration", "LangGraph  StateGraph"],
        ["UI / Dashboard", "Streamlit"],
        ["Databases", "SQLite  (rx_claims.db  +  reports.db)"],
        ["Output Files", "CSV  RCL (Reprocess Claims List) files"],
        ["Language", "Python 3.10+"],
    ],
    col_widths=[2.5, 4.5],
)

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 2. SYSTEM ARCHITECTURE
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "2.  System Architecture", level=1)

add_heading(doc, "High-Level Architecture", level=2)
add_para(doc,
    "The LangGraph pipeline is a single-node graph. All routing, agent invocation, and state management "
    "happen dynamically inside the orchestrator's ReAct loop — driven entirely by the LLM reading "
    "the SOP for each claim.",
    size=10.5)
doc.add_paragraph()

add_code_block(doc,
"""┌──────────────────────────────────────────────────────────────────────┐
│                        LANGGRAPH PIPELINE                            │
│                                                                      │
│   INPUT                                                              │
│  (claim_id,  ──►   ORCHESTRATOR NODE   ──────────────────►  END     │
│   error_code)      [Dynamic ReAct Agent]                             │
│                                                                      │
│             Reads SOP ──► Fetches Claim ──► Calls Agents as Tools   │
└──────────────────────────────────────────────────────────────────────┘""")

doc.add_paragraph()
add_heading(doc, "Full Multi-Agent Component Diagram", level=2)
add_code_block(doc,
"""                    ┌─────────────────────────────────┐
                    │     STREAMLIT UI  (app.py)       │
                    │  Dashboard · Process · Reports   │
                    └────────────────┬────────────────-┘
                                     │ invoke
                                     ▼
               ┌──────────────────────────────────────────┐
               │          LANGGRAPH  (graph.py)           │
               │   StateGraph:  ORCHESTRATOR  ──►  END   │
               └────────────────────┬─────────────────────┘
                                     │
                                     ▼
      ┌────────────────────────────────────────────────────────────┐
      │          DYNAMIC ORCHESTRATOR  (orchestrator.py)           │
      │  AzureChatOpenAI · create_agent() · ReAct Loop             │
      │                                                            │
      │  Shared context bag: _ctx  (reset each run)               │
      │                                                            │
      │  tool_run_rx_agent  ──────────────► RX_CLAIM_AGENT         │
      │  tool_run_report_agent ──────────► REPORT_BUILDER_AGENT    │
      │  tool_run_servicenow_agent ──────► SERVICENOW_AGENT        │
      │  tool_run_email_agent ───────────► EMAIL_AGENT             │
      └────────────────────────────────────────────────────────────┘
               │              │              │             │
               ▼              ▼              ▼             ▼
      ┌──────────────┐ ┌──────────────┐ ┌──────────┐ ┌──────────┐
      │  RX Claim    │ │    Report    │ │ServiceNow│ │  Email   │
      │   Agent      │ │   Builder   │ │  Agent   │ │  Agent   │
      │              │ │   Agent     │ │          │ │          │
      │ 781: provider│ │ RCL / report│ │ Ticket   │ │Compose + │
      │ 935: dates   │ │ generation  │ │ creation │ │  Send    │
      └──────┬───────┘ └──────┬───────┘ └────┬─────┘ └────┬────-┘
             │                │              │            │
             └────────────────┴──────────────┴────────────┘
                                     │
                                     ▼
      ┌────────────────────────────────────────────────────────────┐
      │                       TOOLS LAYER                          │
      │   tools/db_tools.py  ·  tools/sop_tools.py                │
      │   tools/llm_tools.py                                       │
      └────────────────────────────────────────────────────────────┘
                    │                        │
                    ▼                        ▼
          ┌──────────────────┐    ┌──────────────────┐
          │  rx_claims.db    │    │   reports.db     │
          │   (SQLite)       │    │   (SQLite)       │
          └──────────────────┘    └──────────────────┘""")

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 3. DYNAMIC ORCHESTRATION FLOW
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "3.  Dynamic Orchestration Flow", level=1)

add_para(doc,
    "The orchestrator executes a fixed sequence of up to 6 steps, but dynamically decides "
    "which steps to skip based on the SOP rules it reads and the decision produced by the RX Agent.",
    size=10.5)
doc.add_paragraph()

add_code_block(doc,
"""┌───────────────────────────────────────────────────────────────────────────┐
│               DYNAMIC ORCHESTRATOR  —  Step-by-Step ReAct Loop            │
│                                                                           │
│  STEP 1:  tool_load_sop( error_code )                                     │
│            └─► Load SOP_PDE_781.txt  OR  SOP_PDE_935.txt                 │
│            └─► LLM reads the resolution rules before doing anything       │
│                                                                           │
│  STEP 2:  tool_get_claim( claim_id )                                      │
│            └─► Fetch claim from rx_claims.db                              │
│            └─► Reveals: provider_id, received_date, adjudication_ts       │
│                                                                           │
│  STEP 3:  tool_run_rx_agent( claim_id, error_code )         [ALWAYS]      │
│            └─► Adjudicates the claim                                      │
│            └─► Returns DECISION: REPROCESS | REJECT | ALREADY_PROCESSED  │
│                                                                           │
│  STEP 4:  tool_run_report_agent( claim_id, error_code, decision )         │
│            └─► SKIP if decision = REJECT                                  │
│            └─► REPROCESS       → generates RCL CSV file                  │
│            └─► ALREADY_PROCESSED → saves compliance report               │
│                                                                           │
│  STEP 5:  tool_run_servicenow_agent( claim_id, error_code, rcl_file )     │
│            └─► ONLY if decision = REPROCESS                               │
│            └─► Raises incident ticket in ServiceNow                       │
│                                                                           │
│  STEP 6:  tool_run_email_agent( claim_id, error_code, decision )          │
│            └─► SKIP if decision = REJECT                                  │
│            └─► Sends notification to compliance-team@healthplan.com       │
│                                                                           │
│  FINAL:   Merge _ctx results into LangGraph ClaimState                   │
│            └─► Returns unified state to Streamlit / caller               │
└───────────────────────────────────────────────────────────────────────────┘""")

doc.add_paragraph()

add_heading(doc, "Shared Context Bag (_ctx)", level=2)
add_para(doc,
    "The orchestrator maintains a module-level dictionary _ctx that is reset at the start of each run. "
    "All tool wrappers read from and write into this dict, enabling agent-to-agent result sharing "
    "without touching LangGraph state directly during the ReAct loop.",
    size=10.5)

add_code_block(doc,
"""_ctx = {
  # Set by orchestrator before ReAct loop
  "claim_id":    "C0009",
  "error_code":  "781",
  "sop_text":    "<contents of SOP_PDE_781.txt>",

  # Written by tool_run_rx_agent
  "decision":          "REPROCESS",
  "reason":            "Provider P002 resolved from CMS mapping",
  "provider_id":       "P002",
  "resolved_provider": { "new_provider_id": "P002", "provider_name": "...", "npi": "..." },

  # Written by tool_run_report_agent
  "report":            "<compliance report text>",
  "rcl_file":          "output/RCL_20260303_130143.csv",

  # Written by tool_run_servicenow_agent
  "servicenow_ticket": "INC20260303130143",
  "servicenow_summary": "PDE Reprocessing Required — 3 claims",

  # Written by tool_run_email_agent
  "email_status":  "SENT",
  "email_summary": "Notification sent to compliance team",
}""")

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 4. AGENT DESCRIPTIONS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "4.  Agent Descriptions", level=1)

# ── Orchestrator ──
add_heading(doc, "4.1  Dynamic Orchestrator  (agents/orchestrator.py)", level=2)
add_info_box(doc, "Role: Central coordinator — reads SOP, fetches claim, decides and calls specialist agents as tools.")

add_para(doc, "Type:  LangChain create_agent() ReAct agent (AzureChatOpenAI GPT-4o)", size=10.5)
add_para(doc, "Invoked by:  LangGraph graph node ORCHESTRATOR", size=10.5)

doc.add_paragraph()
add_para(doc, "Responsibilities:", bold=True, size=10.5)
for item in [
    "Load the SOP document for the given error code",
    "Fetch the claim from the database",
    "Reason about which specialist agents are needed and in what order",
    "Call each required specialist agent as a LangChain tool",
    "Propagate results through the shared _ctx context bag",
    "Write the final merged state back to LangGraph ClaimState",
]:
    add_bullet(doc, item)

doc.add_paragraph()
add_para(doc, "Available tools:", bold=True, size=10.5)
add_table(doc,
    headers=["Tool", "Purpose"],
    rows=[
        ["tool_load_sop", "Load the SOP text file for the error code"],
        ["tool_get_claim", "Fetch claim from rx_claims.db"],
        ["tool_run_rx_agent", "Invoke RX Claim adjudication specialist"],
        ["tool_run_report_agent", "Invoke Report Builder specialist"],
        ["tool_run_servicenow_agent", "Invoke ServiceNow ticket specialist"],
        ["tool_run_email_agent", "Invoke Email notification specialist"],
    ],
    col_widths=[2.5, 4.5],
)

doc.add_paragraph()

# ── RX Claim Agent ──
add_heading(doc, "4.2  RX Claim Agent  (agents/rx_claim_agent.py)", level=2)
add_info_box(doc, "Role: Adjudication specialist — handles ALL claim adjudication logic for error codes 781 and 935.")

add_para(doc, "Error 781 — Missing / Invalid Provider ID", bold=True, size=11, color=rgb(0,100,0))
for step in [
    ("1.", "Fetch the claim with tool_get_claim"),
    ("2.", "Look up CMS-validated provider ID in the mapping table with tool_get_provider_mapping"),
    ("3.", "If mapping FOUND  →  update provider ID with tool_update_claim_provider_id"),
    ("4.", "Update status with tool_update_claim_status:"),
    ("",   "  Provider resolved  →  status = READY_FOR_REPROCESS  →  Decision: REPROCESS"),
    ("",   "  No mapping found   →  status = REJECTED              →  Decision: REJECT"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    if step[0]:
        r = p.add_run(f"{step[0]}  {step[1]}")
    else:
        r = p.add_run(step[1])
    r.font.size = Pt(10.5)

doc.add_paragraph()
add_para(doc, "Error 935 — Already Adjudicated", bold=True, size=11, color=rgb(150,70,0))
for step in [
    ("1.", "Fetch the claim with tool_get_claim"),
    ("2.", "Compare received_date vs adjudication_ts with tool_compare_claim_dates"),
    ("3.", "If dates EQUAL       →  status = READY_FOR_REPROCESS  →  Decision: REPROCESS"),
    ("4.", "If adj_ts MORE RECENT →  status = ALREADY_PROCESSED   →  Decision: ALREADY_PROCESSED"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run(f"{step[0]}  {step[1]}")
    r.font.size = Pt(10.5)

doc.add_paragraph()

# ── Report Builder ──
add_heading(doc, "4.3  Report Builder Agent  (agents/report_builder_agent.py)", level=2)
add_info_box(doc, "Role: Compliance reporting — generates RCL files or compliance reports. SKIPPED for REJECT decisions.")

add_para(doc, "REPROCESS path (Error 781 or 935):", bold=True, size=10.5)
for s in [
    "Fetch claim details",
    "Get summary of all READY_FOR_REPROCESS claims",
    "Generate RCL (Reprocess Claims List) CSV file in output/ directory",
    "Save report to reports.db",
]:
    add_bullet(doc, s)

doc.add_paragraph()
add_para(doc, "ALREADY_PROCESSED path (Error 935):", bold=True, size=10.5)
for s in [
    "Fetch claim details",
    "Save compliance report to reports.db",
    "Notes: claim excluded from reprocessing, included in PDE compliance report",
]:
    add_bullet(doc, s)

doc.add_paragraph()

# ── ServiceNow Agent ──
add_heading(doc, "4.4  ServiceNow Agent  (agents/servicenow_agent.py)", level=2)
add_info_box(doc, "Role: Incident management — raises a ServiceNow ticket for REPROCESS decisions ONLY.")

for s in [
    "Calls tool_get_reprocess_claims_summary to get a list of all claims flagged for reprocessing",
    "Creates incident ticket with tool_create_servicenow_ticket",
    "Priority: High (2) if more than 5 claims, Medium (3) otherwise",
    "Category: PDE Reprocessing",
    "Assigned To: PDE Operations Team",
    "In production: would call the ServiceNow Table API  POST /api/now/table/incident",
]:
    add_bullet(doc, s)

doc.add_paragraph()

# ── Email Agent ──
add_heading(doc, "4.5  Email Agent  (agents/email_agent.py)", level=2)
add_info_box(doc, "Role: Notification — composes and sends a professional email after claim processing. SKIPPED for REJECT.")

for s in [
    "Fetches the claim to confirm final status",
    "Composes a 3–5 sentence professional notification email",
    "Sends to: compliance-team@healthplan.com",
    "Uses tool_send_email to simulate the send",
    "In production: would connect to SMTP / SendGrid / Azure Communication Services",
]:
    add_bullet(doc, s)

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 5. DECISION LOGIC & ROUTING TABLE
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "5.  Decision Logic & Routing Table", level=1)

add_heading(doc, "Agent Invocation Matrix", level=2)
add_table(doc,
    headers=["Error Code", "Decision", "RX Agent", "Report", "ServiceNow", "Email"],
    rows=[
        ["781", "REPROCESS",         "✅ Yes", "✅ Yes", "✅ Yes", "✅ Yes"],
        ["781", "REJECT",            "✅ Yes", "❌ Skip", "❌ Skip", "❌ Skip"],
        ["935", "REPROCESS",         "✅ Yes", "✅ Yes", "✅ Yes", "✅ Yes"],
        ["935", "ALREADY_PROCESSED", "✅ Yes", "✅ Yes", "❌ Skip", "✅ Yes"],
    ],
    col_widths=[1.2, 2.0, 1.1, 1.1, 1.2, 1.1],
)

doc.add_paragraph()
add_heading(doc, "Decision Tree", level=2)

add_code_block(doc,
"""Claim Received
      │
      ├──► Error Code 781  (Provider ID)
      │          │
      │          └──► RX Agent: Look up provider mapping table
      │                    │
      │                    ├──► Mapping FOUND  ──────────────────────► REPROCESS
      │                    │    Update provider_id in DB               RX → Report → ServiceNow → Email
      │                    │    Status = READY_FOR_REPROCESS
      │                    │
      │                    └──► Mapping NOT FOUND  ──────────────────► REJECT
      │                         Status = REJECTED                      RX only  (no report / email)
      │
      └──► Error Code 935  (Already Adjudicated)
                 │
                 └──► RX Agent: Compare received_date vs adjudication_ts
                           │
                           ├──► received_date == adjudication_ts  ──► REPROCESS
                           │    Potential duplicate submission         RX → Report → ServiceNow → Email
                           │    Status = READY_FOR_REPROCESS
                           │
                           └──► adjudication_ts > received_date  ──► ALREADY_PROCESSED
                                Claim genuinely processed after         RX → Report → Email
                                receipt  (no ServiceNow ticket)         (NO ServiceNow)""")

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 6. SCENARIO WALKTHROUGHS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "6.  Scenario Walkthroughs", level=1)

# ── Scenario 1 ──
add_heading(doc, "Scenario 1 — Error 781, Missing Provider → REPROCESS", level=2)
add_table(doc,
    headers=["Field", "Value"],
    rows=[
        ["Claim ID", "C0009"],
        ["Error Code", "781"],
        ["Provider ID (initial)", "(blank — missing)"],
        ["Expected Decision", "REPROCESS"],
        ["Agents Invoked", "RX_AGENT → REPORT → SERVICENOW → EMAIL"],
    ],
    col_widths=[2.0, 5.0],
)
doc.add_paragraph()
add_para(doc, "Step-by-step execution:", bold=True, size=10.5)
for step in [
    "Orchestrator reads SOP 781: must validate provider ID and resolve from CMS mapping.",
    "RX Agent calls tool_get_claim('C0009')  →  provider_id is blank.",
    "RX Agent calls tool_get_provider_mapping('')  →  finds mapping  →  returns P002.",
    "RX Agent calls tool_update_claim_provider_id('C0009', 'P002')  →  updates DB.",
    "RX Agent calls tool_update_claim_status('C0009', 'READY_FOR_REPROCESS').",
    "Decision = REPROCESS.",
    "Report Agent generates RCL CSV file, saves report to reports.db.",
    "ServiceNow Agent creates incident ticket INC<timestamp>, priority Medium/High.",
    "Email Agent sends notification to compliance-team@healthplan.com.",
]:
    add_bullet(doc, step)

doc.add_paragraph()

# ── Scenario 2 ──
add_heading(doc, "Scenario 2 — Error 935, Equal Dates → REPROCESS", level=2)
add_table(doc,
    headers=["Field", "Value"],
    rows=[
        ["Claim ID", "C0010"],
        ["Error Code", "935"],
        ["received_date", "2025-01-02"],
        ["adjudication_ts", "2025-01-02  (equal to received_date)"],
        ["Expected Decision", "REPROCESS"],
        ["Agents Invoked", "RX_AGENT → REPORT → SERVICENOW → EMAIL"],
    ],
    col_widths=[2.0, 5.0],
)
doc.add_paragraph()
add_para(doc, "Step-by-step execution:", bold=True, size=10.5)
for step in [
    "Orchestrator reads SOP 935: compare dates to determine if claim is a duplicate.",
    "RX Agent calls tool_get_claim('C0010')  →  both dates are 2025-01-02.",
    "RX Agent calls tool_compare_claim_dates('C0010')  →  EQUAL → REPROCESS.",
    "RX Agent calls tool_update_claim_status('C0010', 'READY_FOR_REPROCESS').",
    "Decision = REPROCESS.",
    "Report Agent generates RCL CSV file.",
    "ServiceNow Agent creates reprocessing incident ticket.",
    "Email Agent sends notification.",
]:
    add_bullet(doc, step)

doc.add_paragraph()

# ── Scenario 3 ──
add_heading(doc, "Scenario 3 — Error 935, Adjudication Newer → ALREADY_PROCESSED", level=2)
add_table(doc,
    headers=["Field", "Value"],
    rows=[
        ["Claim ID", "C0006"],
        ["Error Code", "935"],
        ["received_date", "2024-12-15"],
        ["adjudication_ts", "2025-01-02  (MORE RECENT than received_date)"],
        ["Expected Decision", "ALREADY_PROCESSED"],
        ["Agents Invoked", "RX_AGENT → REPORT → EMAIL  (ServiceNow SKIPPED)"],
    ],
    col_widths=[2.0, 5.0],
)
doc.add_paragraph()
add_para(doc, "Step-by-step execution:", bold=True, size=10.5)
for step in [
    "Orchestrator reads SOP 935: compare dates.",
    "RX Agent calls tool_compare_claim_dates('C0006')  →  ADJUDICATION_MORE_RECENT → ALREADY_PROCESSED.",
    "RX Agent calls tool_update_claim_status('C0006', 'ALREADY_PROCESSED').",
    "Decision = ALREADY_PROCESSED.",
    "Report Agent saves compliance report (no RCL file generated).",
    "ServiceNow Agent:  ⏭ SKIPPED  — SOP says do not reprocess, no incident needed.",
    "Email Agent sends compliance notification.",
]:
    add_bullet(doc, step)

doc.add_paragraph()

# ── Scenario 4 ──
add_heading(doc, "Scenario 4 — Error 781, No Provider Mapping → REJECT", level=2)
add_table(doc,
    headers=["Field", "Value"],
    rows=[
        ["Claim ID", "TEST-REJECT-001"],
        ["Error Code", "781"],
        ["Provider ID (initial)", "P999  (invalid — no mapping in DB)"],
        ["Expected Decision", "REJECT"],
        ["Agents Invoked", "RX_AGENT only  (Report, ServiceNow, Email all SKIPPED)"],
    ],
    col_widths=[2.0, 5.0],
)
doc.add_paragraph()
add_para(doc, "Step-by-step execution:", bold=True, size=10.5)
for step in [
    "Orchestrator reads SOP 781: must validate provider ID.",
    "RX Agent calls tool_get_claim('TEST-REJECT-001')  →  provider_id = P999.",
    "RX Agent calls tool_get_provider_mapping('P999')  →  no mapping found.",
    "RX Agent calls tool_update_claim_status('TEST-REJECT-001', 'REJECTED').",
    "Decision = REJECT.",
    "Report Agent:     ⏭ SKIPPED  — no report for rejected claims.",
    "ServiceNow Agent: ⏭ SKIPPED  — no ticket for rejected claims.",
    "Email Agent:      ⏭ SKIPPED  — no notification for rejected claims.",
]:
    add_bullet(doc, step)

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 7. DATABASE SCHEMA
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "7.  Database Schema", level=1)

add_heading(doc, "db/rx_claims.db", level=2)

add_heading(doc, "Table: claims", level=3)
add_table(doc,
    headers=["Column", "Type", "Description"],
    rows=[
        ["claim_id",       "TEXT (PK)", "Unique claim identifier (e.g. C0001, TEST-REJECT-001)"],
        ["error_code",     "TEXT",      "CMS error code: 781 or 935"],
        ["provider_id",    "TEXT",      "Provider NPI / ID (may be blank for 781 errors)"],
        ["adjudication_ts","TEXT",      "Date claim was adjudicated (YYYY-MM-DD)"],
        ["received_date",  "TEXT",      "Date claim was received (YYYY-MM-DD)"],
        ["status",         "TEXT",      "Processing status (see below)"],
    ],
    col_widths=[1.6, 1.3, 4.0],
)
doc.add_paragraph()
add_para(doc, "Valid status values:", bold=True, size=10.5)
add_table(doc,
    headers=["Status", "Meaning"],
    rows=[
        ["NEW",                  "Claim received, not yet processed"],
        ["PENDING",              "Processing in progress"],
        ["READY_FOR_REPROCESS",  "Resolved, queued for CMS resubmission"],
        ["ALREADY_PROCESSED",    "Claim was already adjudicated, excluded from reprocessing"],
        ["REJECTED",             "Claim could not be resolved, permanently rejected"],
    ],
    col_widths=[2.2, 4.8],
)

doc.add_paragraph()
add_heading(doc, "Table: provider_mapping", level=3)
add_table(doc,
    headers=["Column", "Type", "Description"],
    rows=[
        ["old_provider_id", "TEXT", "Invalid / missing provider ID  (empty string = missing)"],
        ["new_provider_id", "TEXT", "Correct CMS-validated provider ID"],
        ["provider_name",   "TEXT", "Provider organisation name"],
        ["npi",             "TEXT", "National Provider Identifier"],
    ],
    col_widths=[1.6, 1.3, 4.0],
)

doc.add_paragraph()
add_heading(doc, "db/reports.db", level=2)
add_heading(doc, "Table: reports", level=3)
add_table(doc,
    headers=["Column", "Type", "Description"],
    rows=[
        ["claim_id",   "TEXT", "Reference to the processed claim"],
        ["error_code", "TEXT", "781 or 935"],
        ["provider_id","TEXT", "Provider ID at time of processing"],
        ["decision",   "TEXT", "REPROCESS, REJECT, or ALREADY_PROCESSED"],
        ["reason",     "TEXT", "Detailed explanation of the decision"],
        ["created_ts", "TEXT", "ISO timestamp of report creation"],
    ],
    col_widths=[1.6, 1.3, 4.0],
)

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 8. TOOLS REFERENCE
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "8.  Tools Reference", level=1)

add_heading(doc, "tools/db_tools.py  —  Database Tools", level=2)
add_table(doc,
    headers=["Tool Function", "Parameters", "Description"],
    rows=[
        ["tool_get_claim",               "claim_id",                              "Fetch full claim record from rx_claims.db"],
        ["tool_get_provider_mapping",    "old_provider_id",                       "Look up CMS-validated provider from mapping table"],
        ["tool_update_claim_provider_id","claim_id, new_provider_id",             "Update provider ID on a claim record"],
        ["tool_update_claim_status",     "claim_id, status",                      "Update processing status of a claim"],
        ["tool_compare_claim_dates",     "claim_id",                              "Compare received_date vs adjudication_ts (935 claims)"],
        ["tool_insert_report",           "claim_id, error_code, provider_id, decision, reason", "Save processed report to reports.db"],
        ["tool_generate_rcl_file",       "(none)",                                "Generate RCL CSV file from all READY_FOR_REPROCESS claims"],
        ["tool_get_reprocess_claims_summary", "(none)",                           "List all claims currently marked READY_FOR_REPROCESS"],
    ],
    col_widths=[2.2, 2.2, 2.6],
)

doc.add_paragraph()
add_heading(doc, "tools/sop_tools.py  —  SOP Tools", level=2)
add_table(doc,
    headers=["Tool Function", "Parameters", "Description"],
    rows=[
        ["tool_load_sop", "error_code", "Load the SOP text file for a given error code (781 or 935)"],
    ],
    col_widths=[2.2, 1.5, 3.3],
)

doc.add_paragraph()
add_heading(doc, "tools/llm_tools.py  —  LLM Instance", level=2)
add_table(doc,
    headers=["Export", "Description"],
    rows=[
        ["llm",         "Shared AzureChatOpenAI instance used by all ReAct agents"],
        ["call_llm(prompt)", "One-shot LLM call helper (returns response string)"],
    ],
    col_widths=[2.5, 4.5],
)

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 9. FILE STRUCTURE
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "9.  File Structure", level=1)

add_code_block(doc,
"""Multi Agentic POC/
│
├── app.py                          Streamlit dashboard (UI)
├── graph.py                        LangGraph pipeline  (ClaimState + build_graph)
├── run.py                          CLI runner (non-Streamlit)
│
├── agents/
│   ├── orchestrator.py             ★ Dynamic ReAct orchestrator (main coordinator)
│   ├── rx_claim_agent.py           ★ Adjudication specialist (781 + 935)
│   ├── report_builder_agent.py     ★ Report / RCL file generator
│   ├── servicenow_agent.py         ★ ServiceNow ticket creator
│   ├── email_agent.py              ★ Email notification sender
│   ├── pde_file_reader.py            PDE flat file reader (ingest utility)
│   └── sop_reader.py                 SOP file reader utility
│
├── tools/
│   ├── db_tools.py                 ★ All database LangChain @tool functions
│   ├── sop_tools.py                ★ SOP loading tool
│   └── llm_tools.py                ★ Shared AzureChatOpenAI instance
│
├── sop/
│   ├── SOP_PDE_781.txt               Standard Operating Procedure for Error 781
│   └── SOP_PDE_935.txt               Standard Operating Procedure for Error 935
│
├── db/
│   ├── rx_claims.db                  Claims database (SQLite)
│   └── reports.db                    Reports database (SQLite)
│
├── output/
│   └── RCL_*.csv                     Generated Reprocess Claims List files
│
├── data/
│   └── PDE_flat_file.txt             Sample PDE flat file for ingestion
│
├── orchestrator_test.py            ★ End-to-end test (all 4 scenarios)
├── init_rx_claims_db.py              DB initialisation (claims + provider mapping)
├── init_reports_db.py                DB initialisation (reports)
├── _seed_reject_claim.py             Seed REJECT test claim (TEST-REJECT-001)
├── check_claims.py                   DB inspection utility
├── _check_mapping.py                 Provider mapping inspection utility
├── insert_test_claim.py              Insert test claim utility
├── ingest_pde.py                     PDE flat file ingestion
├── requirements.txt                  Python dependencies
├── README.md                         Project documentation (Markdown)
└── PROJECT_DOCUMENTATION.docx       ★ This document""")

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 10. SETUP & RUNNING
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "10.  Setup & Running the Application", level=1)

add_heading(doc, "Prerequisites", level=2)
for item in [
    "Python 3.10 or later",
    "Azure OpenAI resource with a GPT-4o deployment",
    "A .env file in the project root with credentials (see below)",
]:
    add_bullet(doc, item)

doc.add_paragraph()
add_heading(doc, ".env File", level=2)
add_code_block(doc,
"""AZURE_OPENAI_API_KEY      = your_api_key_here
AZURE_OPENAI_API_VERSION  = 2024-02-01
AZURE_OPENAI_ENDPOINT     = https://your-resource.openai.azure.com/
AZURE_OPENAI_DEPLOYMENT   = gpt-4o""")

doc.add_paragraph()
add_heading(doc, "Install Dependencies", level=2)
add_code_block(doc, "pip install -r requirements.txt")

doc.add_paragraph()
add_heading(doc, "Initialise Databases  (first time only)", level=2)
add_code_block(doc,
"""python init_rx_claims_db.py
python init_reports_db.py
python _seed_reject_claim.py       # seeds TEST-REJECT-001 for scenario 4""")

doc.add_paragraph()
add_heading(doc, "Run the Streamlit Dashboard", level=2)
add_code_block(doc, "streamlit run app.py")
add_para(doc, "Opens at  http://localhost:8501  with four pages:", size=10.5)
for page in [
    "🏠 Dashboard  —  Metrics, recent reports, workflow diagram",
    "🔍 Process Claim  —  Select a claim and run the agent pipeline",
    "📊 Reports  —  Browse all generated compliance reports",
    "🗄️ Database  —  Explore claims, provider mapping, and reports tables",
]:
    add_bullet(doc, page)

doc.add_paragraph()
add_heading(doc, "Run a Claim Programmatically", level=2)
add_code_block(doc,
"""from graph import build_graph

app = build_graph()
result = app.invoke({
    "claim_id":   "C0009",
    "error_code": "781",
})

print(result["decision"])         # REPROCESS
print(result["agents_invoked"])   # ['RX_AGENT', 'REPORT', 'SERVICENOW', 'EMAIL']
print(result["report"])           # compliance report text
print(result["rcl_file"])         # path to generated RCL CSV""")

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 11. TESTING
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "11.  Testing", level=1)

add_heading(doc, "End-to-End Test  (orchestrator_test.py)", level=2)
add_para(doc, "Runs all 4 scenarios against the live database and LLM:", size=10.5)
add_code_block(doc, "python orchestrator_test.py")

doc.add_paragraph()
add_table(doc,
    headers=["#", "Claim", "Error", "Expected Decision", "Expected Agents", "Excluded Agents"],
    rows=[
        ["1", "C0009",           "781", "REPROCESS",         "RX → Report → ServiceNow → Email", "—"],
        ["2", "C0010",           "935", "REPROCESS",         "RX → Report → ServiceNow → Email", "—"],
        ["3", "C0006",           "935", "ALREADY_PROCESSED", "RX → Report → Email",              "ServiceNow"],
        ["4", "TEST-REJECT-001", "781", "REJECT",            "RX only",                          "Report, ServiceNow, Email"],
    ],
    col_widths=[0.3, 1.5, 0.7, 1.7, 2.2, 1.6],
)

doc.add_paragraph()
add_para(doc, "Each test case validates:", bold=True, size=10.5)
for v in [
    "✅  Correct decision returned (REPROCESS / REJECT / ALREADY_PROCESSED)",
    "✅  All expected agents were invoked",
    "✅  No excluded agents were invoked (e.g. ServiceNow not called for ALREADY_PROCESSED)",
]:
    add_bullet(doc, v)

doc.add_paragraph()
add_heading(doc, "Individual Agent Tests", level=2)
add_code_block(doc,
"""python test_rx_claim_agent.py
python test_report_builder_agent.py
python test_llm.py""")

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 12. KEY DESIGN PRINCIPLES
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "12.  Key Design Principles", level=1)

add_table(doc,
    headers=["Principle", "Implementation"],
    rows=[
        ["SOP-Driven",
         "Orchestrator reads the SOP document before making any decision. The SOP defines the rules; the LLM enforces them."],
        ["LLM-Driven Routing",
         "No hardcoded if/else routing logic. The LLM decides which agents to call based on SOP rules and claim data."],
        ["ReAct Pattern",
         "All agents (Orchestrator, RX, Report, ServiceNow, Email) use LangChain's create_agent() with a Reasoning + Acting loop."],
        ["Tool Isolation",
         "Each specialist agent is only given the tools it needs — preventing cross-contamination and reducing LLM hallucination."],
        ["Shared Context Bag",
         "The _ctx dict allows tool wrappers to pass results between agents without modifying LangGraph state mid-loop."],
        ["Single Node Graph",
         "LangGraph has exactly 1 node (ORCHESTRATOR → END). All routing is internal to the LLM — not encoded as graph edges."],
        ["Traceable",
         "Every agent logs its full ReAct trace (tool calls, LLM reasoning). Traces are visible per-agent in the Streamlit UI."],
        ["Idempotent",
         "Each run resets _ctx and re-evaluates from the SOP — prior runs do not contaminate subsequent ones."],
    ],
    col_widths=[2.0, 5.0],
)

doc.add_paragraph()

add_heading(doc, "Architecture Evolution", level=2)
add_para(doc, "Previous architecture (static routing):", bold=True, size=10.5)
add_code_block(doc,
"""ORCHESTRATOR ──► RX_AGENT ──► REPORT ──► SERVICENOW ──► EMAIL ──► END
(hardcoded graph edges, decision-based routing in Python if/else)""")

doc.add_paragraph()
add_para(doc, "Current architecture (dynamic, LLM-driven):", bold=True, size=10.5)
add_code_block(doc,
"""ORCHESTRATOR (runs full pipeline dynamically via LangChain tools) ──► END
(routing decided at runtime by LLM reading SOP per claim)""")

doc.add_paragraph()
add_info_box(doc,
    "The dynamic architecture means adding a new error code or changing the routing logic "
    "only requires updating the SOP text file — no code changes needed.")

# ─────────────────────────────────────────────────────────────────────────────
# FOOTER NOTE
# ─────────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Built with  LangChain · LangGraph · Azure OpenAI GPT-4o · Streamlit · SQLite")
r.font.size = Pt(9)
r.font.color.rgb = rgb(120, 120, 120)
r.italic = True

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
output_path = "PROJECT_DOCUMENTATION.docx"
doc.save(output_path)
print(f"✅  Word document saved: {output_path}")
